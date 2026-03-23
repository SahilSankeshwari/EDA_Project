import io
import os
import re
import uuid
from typing import Dict, List

import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from django.conf import settings
from django.http import FileResponse, Http404, JsonResponse
from rest_framework.parsers import MultiPartParser, FormParser
from django.shortcuts import render
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from docx import Document
from rest_framework.views import APIView

from .models import DatasetAnalysis


def _safe_name(name: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_-]+", "_", name)[:50]


def _numeric_statements(col: str, s: pd.Series) -> List[str]:
    records = int(s.shape[0])
    mean = s.mean()
    median = s.median()
    vmin = s.min()
    vmax = s.max()
    std = s.std()
    return [
        f"{col} contains {records} records.",
        f"Average value is {mean:.4f}.",
        f"Median value is {median:.4f}.",
        f"Minimum value is {vmin:.4f}.",
        f"Maximum value is {vmax:.4f}.",
        f"Standard deviation is {std:.4f}.",
        f"50% of the data lies at or below {median:.4f}.",
    ]


def _object_statements(col: str, s: pd.Series) -> List[str]:
    s = s.dropna().astype(str)
    records = int(s.shape[0])
    uniq = int(s.nunique())
    mode_val = None
    mode_freq = 0
    if records > 0:
        vc = s.value_counts()
        mode_val = vc.index[0]
        mode_freq = int(vc.iloc[0])
    stmts = [f"{col} contains {records} records.", f"{col} has {uniq} unique values."]
    if mode_val is not None:
        stmts.append(f"Most frequent value is '{mode_val}'.")
        stmts.append(f"It appears {mode_freq} times.")
    return stmts


def _iqr_outliers(s: pd.Series) -> Dict[str, int]:
    q1 = s.quantile(0.25)
    q3 = s.quantile(0.75)
    iqr = q3 - q1
    lower = q1 - 1.5 * iqr
    upper = q3 + 1.5 * iqr
    mask = (s < lower) | (s > upper)
    return {"count": int(mask.sum()), "lower": float(lower), "upper": float(upper)}


class UploadAnalyzeView(APIView):
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request):
        try:
            f = request.FILES.get("file")
            if not f:
                return JsonResponse({"error": "No file uploaded"}, status=400)
            name = f.name.lower()
            if not (name.endswith(".csv") or name.endswith(".xlsx")):
                return JsonResponse({"error": "Only CSV and Excel allowed"}, status=400)

            try:
                if name.endswith(".csv"):
                    df = pd.read_csv(f)
                else:
                    df = pd.read_excel(f, engine="openpyxl")
            except Exception:
                return JsonResponse({"error": "Invalid file format"}, status=400)

            if df.empty:
                return JsonResponse({"error": "File is empty"}, status=400)

            # Separate numeric and object columns for correct describe
            numeric_df = df.select_dtypes(include="number")
            object_df = df.select_dtypes(exclude="number")
            numeric_desc = numeric_df.describe().transpose() if not numeric_df.empty else pd.DataFrame()
            object_desc = object_df.describe().transpose() if not object_df.empty else pd.DataFrame()
            # Combine for a single table output (object rows won't show mean/std/etc.)
            if not numeric_desc.empty and not object_desc.empty:
                desc = pd.concat([numeric_desc, object_desc], axis=0)
            elif not numeric_desc.empty:
                desc = numeric_desc
            else:
                desc = object_desc

            numeric_cols = numeric_df.columns.tolist()

            outliers: Dict[str, Dict[str, int]] = {}
            session_id = f"a{uuid.uuid4().hex[:8]}"
            outputs_dir = os.path.join(settings.MEDIA_ROOT, "outputs", session_id)
            graphs_dir = os.path.join(settings.MEDIA_ROOT, "graphs", session_id)
            os.makedirs(outputs_dir, exist_ok=True)
            os.makedirs(graphs_dir, exist_ok=True)
            graph_urls: List[str] = []

            # Only meaningful graphs:
            # 1) Bar chart of mean values for numeric columns
            if not numeric_df.empty:
                try:
                    means = numeric_df.mean().dropna()
                    plt.figure(figsize=(8, 4))
                    means.sort_values(ascending=False).plot(kind="bar", color="#0d6efd")
                    plt.title("Mean Values of Numeric Columns")
                    plt.ylabel("Mean")
                    plt.tight_layout()
                    bar_path = os.path.join(graphs_dir, "means_bar.png")
                    plt.savefig(bar_path)
                    plt.close()
                    graph_urls.append(settings.MEDIA_URL + f"graphs/{session_id}/means_bar.png")
                except Exception:
                    pass
                # 2) Boxplot for all numeric columns
                try:
                    plt.figure(figsize=(8, 4))
                    sns.boxplot(data=numeric_df, orient="h", palette="Set2")
                    plt.title("Boxplot of Numeric Columns")
                    plt.tight_layout()
                    boxplot_path = os.path.join(graphs_dir, "numeric_boxplot.png")
                    plt.savefig(boxplot_path)
                    plt.close()
                    graph_urls.append(settings.MEDIA_URL + f"graphs/{session_id}/numeric_boxplot.png")
                except Exception:
                    pass
                # 3) Correlation heatmap (optional)
                try:
                    if numeric_df.shape[1] >= 2:
                        corr = numeric_df.corr(numeric_only=True)
                        plt.figure(figsize=(6, 5))
                        sns.heatmap(corr, annot=False, cmap="Blues")
                        plt.title("Correlation Heatmap")
                        plt.tight_layout()
                        heatmap_path = os.path.join(graphs_dir, "correlation_heatmap.png")
                        plt.savefig(heatmap_path)
                        plt.close()
                        graph_urls.append(settings.MEDIA_URL + f"graphs/{session_id}/correlation_heatmap.png")
                except Exception:
                    pass

            # Compute outliers using IQR for numeric columns
            for col in numeric_cols:
                s = numeric_df[col].dropna()
                outliers[col] = _iqr_outliers(s)

            preview = df.head(5).fillna("").astype(str).to_dict(orient="records")
            desc_rows = desc.reset_index().fillna("").astype(str).to_dict(orient="records")
            outlier_counts = {c: v["count"] for c, v in outliers.items()}

            # Build statements for numeric and object columns
            numeric_statements: Dict[str, List[str]] = {}
            object_statements: Dict[str, List[str]] = {}
            for col in numeric_cols:
                s = numeric_df[col].dropna()
                numeric_statements[col] = _numeric_statements(col, s)
            for col in object_df.columns.tolist():
                s = object_df[col].dropna()
                object_statements[col] = _object_statements(col, s)

            # Save summary_report.xlsx / .pdf / .docx
            # Excel
            try:
                with pd.ExcelWriter(os.path.join(outputs_dir, "summary_report.xlsx"), engine="openpyxl") as writer:
                    pd.DataFrame(desc_rows).to_excel(writer, sheet_name="Describe", index=False)
                    # Flatten statements to two columns
                    num_rows = [{"column": c, "statement": st} for c, arr in numeric_statements.items() for st in arr]
                    obj_rows = [{"column": c, "statement": st} for c, arr in object_statements.items() for st in arr]
                    if num_rows:
                        pd.DataFrame(num_rows).to_excel(writer, sheet_name="NumericStatements", index=False)
                    if obj_rows:
                        pd.DataFrame(obj_rows).to_excel(writer, sheet_name="ObjectStatements", index=False)
                    if outlier_counts:
                        pd.DataFrame([{"column": c, "outliers": n} for c, n in outlier_counts.items()]).to_excel(
                            writer, sheet_name="Outliers", index=False
                        )
            except Exception:
                pass

            # PDF with simple English explanations
            try:
                pdf_buf = io.BytesIO()
                doc = SimpleDocTemplate(pdf_buf, pagesize=letter)
                styles = getSampleStyleSheet()
                elements = [Paragraph("StatSnap – Smart Data Analyzer Report", styles["Title"]), Spacer(1, 12)]
                elements.append(Paragraph(f"The dataset contains {df.shape[0]} rows and {df.shape[1]} columns.", styles["BodyText"]))
                elements.append(Spacer(1, 6))
                for col, stmts in numeric_statements.items():
                    elements.append(Paragraph(f"<b>{col}</b>", styles["Heading3"]))
                    for s in stmts:
                        elements.append(Paragraph(s, styles["BodyText"]))
                    count = outlier_counts.get(col, 0)
                    elements.append(Paragraph(f"Outliers detected: {count}.", styles["BodyText"]))
                    elements.append(Spacer(1, 6))
                for col, stmts in object_statements.items():
                    elements.append(Paragraph(f"<b>{col}</b>", styles["Heading3"]))
                    for s in stmts:
                        elements.append(Paragraph(s, styles["BodyText"]))
                    elements.append(Spacer(1, 6))
                doc.build(elements)
                default_storage.save(f"outputs/{session_id}/summary_report.pdf", ContentFile(pdf_buf.getvalue()))
            except Exception:
                pass

            # DOCX with simple explanations
            try:
                document = Document()
                document.add_heading("StatSnap – Smart Data Analyzer Report", level=1)
                document.add_paragraph(f"The dataset contains {df.shape[0]} rows and {df.shape[1]} columns.")
                document.add_paragraph("Numeric Columns")
                for col, stmts in numeric_statements.items():
                    document.add_paragraph(col, style="Heading3")
                    for s in stmts:
                        document.add_paragraph(s, style="List Bullet")
                    document.add_paragraph(
                        f"Outliers detected: {outlier_counts.get(col, 0)}.", style="List Bullet"
                    )
                document.add_paragraph("Object Columns")
                for col, stmts in object_statements.items():
                    document.add_paragraph(col, style="Heading3")
                    for s in stmts:
                        document.add_paragraph(s, style="List Bullet")
                docx_bytes = io.BytesIO()
                document.save(docx_bytes)
                default_storage.save(f"outputs/{session_id}/summary_report.docx", ContentFile(docx_bytes.getvalue()))
            except Exception:
                pass

            return JsonResponse(
                {
                    "preview": preview,
                    "describe": desc_rows,
                    "statements": {"numeric": numeric_statements, "object": object_statements},
                    "outliers": outlier_counts,
                    "graphs": graph_urls,
                    "totals": {"rows": int(df.shape[0]), "numeric_cols": len(numeric_cols), "object_cols": len(object_df.columns.tolist())},
                    "downloads": {
                        "excel": "/api/download/excel/",
                        "pdf": "/api/download/pdf/",
                        "docx": "/api/download/docx/",
                    },
                }
            )
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)


def home(request):
    return render(request, "index.html")


class DownloadView(APIView):
    def get(self, request, fmt: str):
        fmt = fmt.lower()
        outputs_root = os.path.join(settings.MEDIA_ROOT, "outputs")
        if not os.path.exists(outputs_root):
            return JsonResponse({"error": "No analysis available"}, status=404)
        try:
            sessions = sorted(
                [os.path.join(outputs_root, d) for d in os.listdir(outputs_root)],
                key=lambda p: os.path.getmtime(p),
                reverse=True,
            )
        except Exception:
            return JsonResponse({"error": "No analysis available"}, status=404)
        name_map = {
            "excel": "summary_report.xlsx",
            "pdf": "summary_report.pdf",
            "docx": "summary_report.docx",
        }
        if fmt not in name_map:
            return JsonResponse({"error": "Unsupported format"}, status=400)
        target_name = name_map[fmt]
        for sess in sessions:
            candidate = os.path.join(sess, target_name)
            if os.path.exists(candidate):
                return FileResponse(open(candidate, "rb"), as_attachment=True, filename=target_name)
        return JsonResponse({"error": "Requested file not found"}, status=404)
